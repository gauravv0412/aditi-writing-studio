"""Writing Studio — a tiny personal tool that drafts blog articles in one
person's voice, with a live rich-text editor, an AI chat that edits the draft,
full version history, Word/Google-Docs export, and dynamic learning.

Run locally:   uvicorn app:app --reload --port 8765
Then open:     http://localhost:8765
"""
import hashlib
import io
import json
import os
import re
import time
from html import unescape
from html.parser import HTMLParser

import anthropic
from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, Depends, FastAPI, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

import db
import prompts
import seed

HERE = os.path.dirname(__file__)
MODEL = os.environ.get("MODEL", "claude-opus-4-8")
STYLE_MODEL = os.environ.get("STYLE_MODEL", "claude-sonnet-4-6")  # fast helper for style-rule extraction
APP_PASSWORD = os.environ.get("APP_PASSWORD", "").strip()
AUTH_TOKEN = hashlib.sha256(f"studio::{APP_PASSWORD}".encode()).hexdigest()
MAX_TOKENS = 4096

app = FastAPI(title="Aditi Srivastava — Writing Studio")

# Make sure the database exists and is seeded with her style DNA on startup.
with db.connect() as _c:
    db.init_db(_c)
    seed.seed_if_empty(_c)


# --------------------------------------------------------------------------
# Anthropic helpers
# --------------------------------------------------------------------------
def get_client(conn):
    key = os.environ.get("ANTHROPIC_API_KEY") or db.get_setting(conn, "api_key")
    if not key:
        return None
    return anthropic.Anthropic(api_key=key)


def sse(obj) -> str:
    return f"data: {json.dumps(obj)}\n\n"


# --------------------------------------------------------------------------
# HTML helpers (content is rich HTML)
# --------------------------------------------------------------------------
_TAG_RE = re.compile(r"<[^>]+>")
_BLOCK_BREAK = re.compile(r"(?i)<\s*(br|/p|/h[1-6]|/div|/li|/tr)\s*/?>")


def clean_html(html: str) -> str:
    """Strip code fences, scripts, and event handlers from model/user HTML."""
    if not html:
        return ""
    h = html.strip()
    h = re.sub(r"^```[a-zA-Z]*\s*", "", h)
    h = re.sub(r"\s*```$", "", h)
    h = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", "", h)
    h = re.sub(r'(?i)\son\w+\s*=\s*"[^"]*"', "", h)
    h = re.sub(r"(?i)\son\w+\s*=\s*'[^']*'", "", h)
    h = re.sub(r"(?i)javascript:", "", h)
    return h.strip()


def html_to_text(html: str) -> str:
    """Plain-text version of HTML, preserving line breaks at block boundaries."""
    s = _BLOCK_BREAK.sub("\n", html or "")
    s = _TAG_RE.sub("", s)
    s = unescape(s)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def first_line_title(content: str, fallback: str) -> str:
    for line in html_to_text(content).splitlines():
        t = line.strip().lstrip("#").strip()
        if t:
            return t[:120]
    return (fallback or "Untitled")[:120]


# ---- HTML -> .docx ---------------------------------------------------------
_FONT_PT = {1: 11, 2: 13, 3: 16, 4: 18, 5: 22, 6: 28, 7: 34}
_HEADING_LEVEL = {"h1": 0, "h2": 1, "h3": 2, "h4": 3, "h5": 4, "h6": 5}
_INLINE_FMT = {"strong", "b", "em", "i", "u", "span", "font"}


class _DocxBuilder(HTMLParser):
    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.para = None
        self.inline = []   # stack of {bold/italic/underline/size}
        self.lists = []     # stack of paragraph styles
        self.skip = 0

    def _fmt(self):
        f = {"bold": False, "italic": False, "underline": False, "size": None}
        for d in self.inline:
            if d.get("bold"): f["bold"] = True
            if d.get("italic"): f["italic"] = True
            if d.get("underline"): f["underline"] = True
            if d.get("size") is not None: f["size"] = d["size"]
        return f

    def _size(self, tag, attrs):
        if tag == "font" and attrs.get("size"):
            try:
                return _FONT_PT.get(int(attrs["size"]))
            except ValueError:
                return None
        m = re.search(r"font-size\s*:\s*([\d.]+)\s*(px|pt|em|rem)?", attrs.get("style", ""), re.I)
        if not m:
            return None
        val = float(m.group(1))
        unit = (m.group(2) or "px").lower()
        if unit == "pt":
            return val
        if unit in ("em", "rem"):
            return round(val * 12, 1)
        return round(val * 0.75, 1)  # px -> pt

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self.skip += 1
            return
        if self.skip:
            return
        a = dict(attrs)
        if tag in _HEADING_LEVEL:
            self.para = self.doc.add_heading("", level=min(_HEADING_LEVEL[tag], 4))
        elif tag in ("p", "div"):
            self.para = self.doc.add_paragraph()
        elif tag == "br":
            if self.para is None:
                self.para = self.doc.add_paragraph()
            self.para.add_run().add_break()
        elif tag in ("ul", "ol"):
            self.lists.append("List Number" if tag == "ol" else "List Bullet")
        elif tag == "li":
            self.para = self.doc.add_paragraph(style=self.lists[-1] if self.lists else None)
        elif tag in ("strong", "b"):
            self.inline.append({"bold": True})
        elif tag in ("em", "i"):
            self.inline.append({"italic": True})
        elif tag == "u":
            self.inline.append({"underline": True})
        elif tag in ("span", "font"):
            self.inline.append({"size": self._size(tag, a)})

    def handle_endtag(self, tag):
        if tag in ("script", "style"):
            self.skip = max(0, self.skip - 1)
            return
        if self.skip:
            return
        if tag in _INLINE_FMT:
            if self.inline:
                self.inline.pop()
        elif tag in ("ul", "ol"):
            if self.lists:
                self.lists.pop()

    def handle_data(self, text):
        if self.skip or not text:
            return
        t = re.sub(r"\s+", " ", text)
        if not t.strip() and self.para is None:
            return
        if self.para is None:
            self.para = self.doc.add_paragraph()
        run = self.para.add_run(t)
        f = self._fmt()
        if f["bold"]:
            run.bold = True
        if f["italic"]:
            run.italic = True
        if f["underline"]:
            run.underline = True
        if f["size"]:
            run.font.size = Pt(f["size"])


def build_docx(title: str, html: str) -> io.BytesIO:
    doc = Document()
    html = html or ""
    if "<h1" not in html.lower() and title:
        doc.add_heading(title, level=0)
    try:
        _DocxBuilder(doc).feed(html)
    except Exception:
        # fall back to plain text if the HTML is unparseable
        for line in html_to_text(html).splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
    if not doc.paragraphs and title:
        doc.add_heading(title, level=0)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------
# Auth
# --------------------------------------------------------------------------
def is_authed(request: Request) -> bool:
    if not APP_PASSWORD:
        return True
    return request.cookies.get("studio_auth") == AUTH_TOKEN


def require_auth(request: Request):
    if not is_authed(request):
        raise HTTPException(status_code=401, detail="Unauthorized")


class LoginBody(BaseModel):
    password: str


@app.post("/api/login")
def login(body: LoginBody):
    if not APP_PASSWORD:
        return JSONResponse({"ok": True})
    if body.password != APP_PASSWORD:
        raise HTTPException(status_code=401, detail="Wrong password")
    resp = JSONResponse({"ok": True})
    resp.set_cookie("studio_auth", AUTH_TOKEN, httponly=True, samesite="lax", max_age=60 * 60 * 24 * 30)
    return resp


@app.get("/api/health")
def health(request: Request):
    with db.connect() as conn:
        key_set = bool(os.environ.get("ANTHROPIC_API_KEY") or db.get_setting(conn, "api_key"))
        n_examples = db.count_style_examples(conn)
    return {
        "password_required": bool(APP_PASSWORD),
        "authed": is_authed(request),
        "api_key_set": key_set,
        "model": MODEL,
        "examples": n_examples,
    }


# --------------------------------------------------------------------------
# Protected API
# --------------------------------------------------------------------------
api = APIRouter(prefix="/api", dependencies=[Depends(require_auth)])


class ApiKeyBody(BaseModel):
    api_key: str


@api.post("/settings/apikey")
def set_api_key(body: ApiKeyBody):
    key = body.api_key.strip()
    with db.connect() as conn:
        db.set_setting(conn, "api_key", key)
    return {"ok": True, "api_key_set": bool(key)}


def article_dict(row):
    return {
        "id": row["id"], "title": row["title"], "topic": row["topic"],
        "status": row["status"], "updated_at": row["updated_at"], "created_at": row["created_at"],
    }


@api.get("/articles")
def get_articles():
    with db.connect() as conn:
        return [article_dict(r) for r in db.list_articles(conn)]


class NewArticleBody(BaseModel):
    title: str = "Untitled"
    topic: str = ""


@api.post("/articles")
def new_article(body: NewArticleBody):
    with db.connect() as conn:
        aid = db.create_article(conn, title=body.title or "Untitled", topic=body.topic)
        return article_dict(db.get_article(conn, aid))


@api.get("/articles/{aid}")
def read_article(aid: int):
    with db.connect() as conn:
        row = db.get_article(conn, aid)
        if not row:
            raise HTTPException(404, "Not found")
        chat = [{"role": m["role"], "content": m["content"], "at": m["created_at"]} for m in db.list_chat(conn, aid)]
        return {
            "id": row["id"], "title": row["title"], "topic": row["topic"],
            "content": row["content"], "status": row["status"], "updated_at": row["updated_at"], "chat": chat,
        }


@api.delete("/articles/{aid}")
def remove_article(aid: int):
    with db.connect() as conn:
        db.delete_article(conn, aid)
    return {"ok": True}


class SaveBody(BaseModel):
    content: str
    title: str | None = None


@api.put("/articles/{aid}")
def save_article(aid: int, body: SaveBody):
    """Manual autosave. Collapses rapid manual saves into one version."""
    with db.connect() as conn:
        row = db.get_article(conn, aid)
        if not row:
            raise HTTPException(404, "Not found")
        content = clean_html(body.content)
        title = body.title if body.title is not None else first_line_title(content, row["topic"])
        changed = content != row["content"]
        db.update_article(conn, aid, content=content, title=title)
        if changed:
            last = conn.execute(
                "SELECT id, source, created_at FROM versions WHERE article_id=? ORDER BY id DESC LIMIT 1",
                (aid,),
            ).fetchone()
            if last and last["source"] == "manual" and (time.time() - last["created_at"]) < 120:
                conn.execute("UPDATE versions SET content=?, created_at=? WHERE id=?",
                             (content, time.time(), last["id"]))
                conn.commit()
            else:
                db.add_version(conn, aid, content, "manual", label="edit")
        return {"ok": True, "title": title}


class GenerateBody(BaseModel):
    topic: str


@api.post("/articles/{aid}/generate")
def generate(aid: int, body: GenerateBody):
    topic = body.topic.strip()

    def stream():
        conn = db.connect()
        try:
            if not db.get_article(conn, aid):
                yield sse({"type": "error", "message": "Article not found"})
                return
            client = get_client(conn)
            if not client:
                yield sse({"type": "error", "message": "No API key set. Open Settings and paste your Anthropic API key."})
                return
            db.update_article(conn, aid, topic=topic)
            system = prompts.build_generation_system(conn)
            user = prompts.build_generation_user(topic)
            chunks = []
            try:
                with client.messages.stream(
                    model=MODEL, max_tokens=MAX_TOKENS, system=system,
                    messages=[{"role": "user", "content": user}],
                ) as s:
                    for text in s.text_stream:
                        chunks.append(text)
                        yield sse({"type": "delta", "text": text})
            except anthropic.APIStatusError as e:
                yield sse({"type": "error", "message": f"API error: {getattr(e, 'message', str(e))}"})
                return
            except Exception as e:  # noqa
                yield sse({"type": "error", "message": str(e)})
                return
            content = clean_html("".join(chunks))
            title = first_line_title(content, topic)
            db.update_article(conn, aid, content=content, title=title)
            vid = db.add_version(conn, aid, content, "generated", label=topic[:60])
            yield sse({"type": "done", "version_id": vid, "title": title, "content": content})
        finally:
            conn.close()

    return StreamingResponse(stream(), media_type="text/event-stream")


STYLE_EXTRACT_SCHEMA = {
    "type": "object", "additionalProperties": False,
    "properties": {
        "pure_style": {"type": "boolean"},
        "add": {"type": "array", "items": {"type": "string"}},
        "remove_ids": {"type": "array", "items": {"type": "integer"}},
        "summary": {"type": "string"},
    },
    "required": ["pure_style", "add", "remove_ids", "summary"],
}


def extract_style_change(conn, client, instruction):
    """Detect durable style-rule changes in a chat message (add / modify / remove)."""
    try:
        msg = client.messages.create(
            model=STYLE_MODEL, max_tokens=700,
            system=prompts.STYLE_EXTRACT_SYSTEM,
            messages=[{"role": "user", "content": prompts.build_style_extract_user(conn, instruction)}],
            output_config={"format": {"type": "json_schema", "schema": STYLE_EXTRACT_SCHEMA}},
        )
        data = json.loads("".join(b.text for b in msg.content if b.type == "text").strip())
        return {
            "pure_style": bool(data.get("pure_style")),
            "add": [s.strip() for s in data.get("add", []) if isinstance(s, str) and s.strip()][:6],
            "remove_ids": [int(i) for i in data.get("remove_ids", []) if str(i).lstrip("-").isdigit()][:12],
            "summary": (data.get("summary") or "").strip(),
        }
    except Exception:
        return {"pure_style": False, "add": [], "remove_ids": [], "summary": ""}


def apply_style_change(conn, change):
    existing = {n["id"] for n in db.list_style_notes(conn)}
    removed = 0
    for rid in dict.fromkeys(change["remove_ids"]):  # dedupe; only count rules that actually exist
        if rid in existing:
            db.delete_style_note(conn, rid)
            removed += 1
    added = 0
    for rule in change["add"]:
        before = len(db.list_style_notes(conn))
        db.add_style_note(conn, rule, source="chat")
        if len(db.list_style_notes(conn)) > before:
            added += 1
    if not added and not removed:
        return None
    if change["summary"]:
        return "📝 " + change["summary"]
    bits = []
    if added:
        bits.append(f"added {added} rule(s)")
    if removed:
        bits.append(f"removed {removed} rule(s)")
    return "📝 Style updated — " + " and ".join(bits) + "."


class ChatBody(BaseModel):
    message: str
    content: str


@api.post("/articles/{aid}/chat")
def chat(aid: int, body: ChatBody):
    instruction = body.message.strip()
    current = body.content

    def stream():
        conn = db.connect()
        try:
            if not db.get_article(conn, aid):
                yield sse({"type": "error", "message": "Article not found"})
                return
            client = get_client(conn)
            if not client:
                yield sse({"type": "error", "message": "No API key set. Open Settings and paste your Anthropic API key."})
                return
            db.add_chat(conn, aid, "user", instruction)

            # First: does this instruction change her durable style rules?
            change = extract_style_change(conn, client, instruction)
            applied = apply_style_change(conn, change)

            # Pure style management — update the rules, leave the article alone.
            if change["pure_style"]:
                summary = applied or "Noted — nothing needed changing in your rules."
                db.add_chat(conn, aid, "assistant", summary)
                yield sse({"type": "style", "summary": summary})
                return

            # Otherwise edit the draft (any new rules are already in effect below).
            system = prompts.build_chat_system(conn)
            user = prompts.build_chat_user(current, instruction)
            chunks = []
            try:
                with client.messages.stream(
                    model=MODEL, max_tokens=MAX_TOKENS, system=system,
                    messages=[{"role": "user", "content": user}],
                ) as s:
                    for text in s.text_stream:
                        chunks.append(text)
                        yield sse({"type": "delta", "text": text})
            except anthropic.APIStatusError as e:
                yield sse({"type": "error", "message": f"API error: {getattr(e, 'message', str(e))}"})
                return
            except Exception as e:  # noqa
                yield sse({"type": "error", "message": str(e)})
                return
            content = clean_html("".join(chunks))
            title = first_line_title(content, db.get_article(conn, aid)["topic"])
            db.update_article(conn, aid, content=content, title=title)
            vid = db.add_version(conn, aid, content, "ai_edit", label=instruction[:60])
            db.add_chat(conn, aid, "assistant", "✓ Updated the draft." + (("  " + applied) if applied else ""))
            done = {"type": "done", "version_id": vid, "title": title, "content": content}
            if applied:
                done["learned"] = applied
            yield sse(done)
        finally:
            conn.close()

    return StreamingResponse(stream(), media_type="text/event-stream")


@api.get("/articles/{aid}/versions")
def versions(aid: int):
    with db.connect() as conn:
        return [
            {"id": v["id"], "source": v["source"], "label": v["label"], "at": v["created_at"], "chars": v["chars"]}
            for v in db.list_versions(conn, aid)
        ]


@api.post("/articles/{aid}/versions/{vid}/restore")
def restore_version(aid: int, vid: int):
    with db.connect() as conn:
        v = db.get_version(conn, vid)
        if not v or v["article_id"] != aid:
            raise HTTPException(404, "Version not found")
        db.update_article(conn, aid, content=v["content"])
        db.add_version(conn, aid, v["content"], "restore", label=f"from #{vid}")
        return {"ok": True, "content": v["content"]}


@api.post("/articles/{aid}/finalize")
def finalize(aid: int):
    """Mark done -> add to the style corpus (as clean text) so the tool keeps learning."""
    with db.connect() as conn:
        row = db.get_article(conn, aid)
        if not row:
            raise HTTPException(404, "Not found")
        plain = html_to_text(row["content"] or "")
        if not plain.strip():
            raise HTTPException(400, "Nothing to learn from — the draft is empty.")
        db.update_article(conn, aid, status="published")
        db.add_style_example(conn, row["title"], plain, source="written")
        refreshed = _refresh_profile(conn)
        return {"ok": True, "examples": db.count_style_examples(conn), "profile_refreshed": refreshed}


def _refresh_profile(conn) -> bool:
    client = get_client(conn)
    if not client:
        return False
    try:
        user = prompts.build_profile_refresh_messages(conn)
        msg = client.messages.create(model=MODEL, max_tokens=MAX_TOKENS, messages=[{"role": "user", "content": user}])
        text = "".join(b.text for b in msg.content if b.type == "text").strip()
        if text:
            db.set_setting(conn, "style_profile", text)
            return True
    except Exception:
        return False
    return False


@api.post("/style/refresh")
def style_refresh():
    with db.connect() as conn:
        ok = _refresh_profile(conn)
        return {"ok": ok, "profile": prompts.get_profile(conn)}


@api.get("/style")
def get_style():
    with db.connect() as conn:
        notes = [{"id": n["id"], "note": n["note"], "source": n["source"]} for n in db.list_style_notes(conn)]
        examples = [{"id": e["id"], "title": e["title"], "source": e["source"]} for e in db.list_style_examples(conn)]
        return {"profile": prompts.get_profile(conn), "notes": notes, "examples": examples}


class NoteBody(BaseModel):
    note: str


@api.post("/style/notes")
def add_note(body: NoteBody):
    with db.connect() as conn:
        db.add_style_note(conn, body.note, source="manual")
    return {"ok": True}


@api.delete("/style/notes/{nid}")
def del_note(nid: int):
    with db.connect() as conn:
        db.delete_style_note(conn, nid)
    return {"ok": True}


@api.get("/articles/{aid}/export.docx")
def export_docx(aid: int):
    with db.connect() as conn:
        row = db.get_article(conn, aid)
        if not row:
            raise HTTPException(404, "Not found")
        buf = build_docx(row["title"], row["content"] or "")
        safe = re.sub(r"[^\w\- ]+", "", (row["title"] or "article")).strip() or "article"
    headers = {"Content-Disposition": f'attachment; filename="{safe}.docx"'}
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


app.include_router(api)

# Static frontend
app.mount("/static", StaticFiles(directory=os.path.join(HERE, "static")), name="static")


@app.get("/")
def index():
    return FileResponse(os.path.join(HERE, "static", "index.html"))
